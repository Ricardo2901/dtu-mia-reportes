from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt
from io import BytesIO

def generar_word(datos_usuario):
    doc = Document()

    # Título con formato
    heading = doc.add_heading(level=1)
    run = heading.add_run("Datos de Usuario")
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 0, 255)

    # Agregamos datos dinámicos
    doc.add_paragraph(f"Nombre: {datos_usuario.get('nombre', '')}")
    doc.add_paragraph(f"Correo: {datos_usuario.get('correo', '')}")
    doc.add_paragraph(f"Edad: {datos_usuario.get('edad', '')}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
